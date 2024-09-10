import importlib.util
import os
import tkinter as tk
from tkinterdnd2 import TkinterDnD, DND_FILES
from tkinter import ttk, filedialog, messagebox
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import asyncio
import threading
import time
from prettytable import PrettyTable
from pathlib import Path
import logging
import subprocess
import sys
import importlib
import configparser
import json

logging.basicConfig(level=logging.INFO, filename='pdf_magic.log', filemode='a', format='%(asctime)s - %(levelname)s - %(message)s')

class PDFMagicApp:

# ------ GUI Setup ------

    def __init__(self):
        self.root = TkinterDnD.Tk()
        self.root.configure(bg='#1a1a1a')
        self.root.title("PDF Magic")
        self.root.geometry("900x700")
        self.after_id = None
        self._setup_ui()
        self._initialize_conversion_table()
        self.progress_var = tk.DoubleVar()
        self._create_progress_bar()
        self.config = configparser.ConfigParser()
        self.config_file = 'pdf_magic_config.ini'
        self.load_settings()
        self._create_settings_ui()

    def _setup_ui(self):
        self._create_file_selection_area()
        self._create_browse_button()
        self._create_drag_and_drop_area()
        self._create_conversion_button()
        self._create_log_text_area()

    def _create_file_selection_area(self):
        # Label und Eingabefeld für Datei
        file_label = tk.Label(self.root, text="PDF-Dateien auswählen: \nNutzen Sie alternativ auch Drag and Drop",
                              bg='#1a1a1a', fg='#ffffff', font=('Helvetica', 16))
        file_label.pack(pady=15)
        
        self.file_entry = tk.Entry(self.root, width=70, bg='#2c2c2c', fg='#ffffff')
        self.file_entry.pack(padx=15, pady=5)

    def load_settings(self):
        self.config.read(self.config_file)
        if 'Settings' not in self.config:
            self.config['Settings'] = {'output_directory': ''}

    def save_settings(self):
        with open(self.config_file, 'w') as configfile:
            self.config.write(configfile)

    def _create_settings_ui(self):
        settings_frame = tk.Frame(self.root, bg='#1a1a1a')
        settings_frame.pack(fill=tk.X, padx=15, pady=5)

        output_dir_label = tk.Label(settings_frame, text="Bevorzugter Ausgabeordner:", bg='#1a1a1a', fg='#ffffff')
        output_dir_label.pack(side=tk.LEFT, padx=5)

        self.output_dir_entry = tk.Entry(settings_frame, width=50, bg='#2c2c2c', fg='#ffffff')
        self.output_dir_entry.pack(side=tk.LEFT, padx=5)
        self.output_dir_entry.insert(0, self.config['Settings']['output_directory'])

        browse_button = tk.Button(settings_frame, text="Durchsuchen", command=self.browse_output_dir, bg='#333333', fg='#ffffff')
        browse_button.pack(side=tk.LEFT, padx=5)

        save_button = tk.Button(settings_frame, text="Einstellungen speichern", command=self.save_output_dir, bg='#333333', fg='#ffffff')
        save_button.pack(side=tk.LEFT, padx=5)

    def browse_output_dir(self):
        directory = filedialog.askdirectory()
        if directory:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, directory)

    def save_output_dir(self):
        output_dir = self.output_dir_entry.get()
        self.config['Settings']['output_directory'] = output_dir
        self.save_settings()
        messagebox.showinfo("Einstellungen gespeichert", "Der bevorzugte Ausgabeordner wurde gespeichert.")

    async def _create_output_directory(self, pdf_path, conversion_type):
        preferred_output_dir = self.config['Settings']['output_directory']
        if preferred_output_dir and os.path.isdir(preferred_output_dir):
            output_dir = os.path.join(preferred_output_dir, f"Konvertierte_{conversion_type.upper()}")
        else:
            output_dir = os.path.join(os.path.dirname(pdf_path), f"Konvertierte_{conversion_type.upper()}")
        
        os.makedirs(output_dir, exist_ok=True)
        return output_dir

    def _create_drag_and_drop_area(self):
        # Globales Drag & Drop implementieren
        self.file_entry.drop_target_register(DND_FILES)  
        self.file_entry.dnd_bind('<<Drop>>', self.drop)

    def _create_conversion_button(self):
        # Button zum Starten der Konvertierung
        convert_button = tk.Button(self.root, text="Konvertieren", 
                                   bg='#444444', fg='#ffffff',
                                   command=self.start_conversion_wrapper,
                                   relief='flat',
                                   overrelief='ridge',
                                   borderwidth=0)
        convert_button.pack(pady=22)

    def start_conversion_wrapper(self):
        threading.Thread(target=self.run_async_conversion, daemon=True).start()

    def run_async_conversion(self):
        asyncio.run(self.start_conversion())

    def _create_browse_button(self):
        # Button zum Auswählen der Datei
        browse_button = tk.Button(self.root, text="Durchsuchen", 
                                  bg='#333333', fg='#ffffff',
                                  command=self.select_file,
                                  relief='flat',
                                  overrelief='ridge',
                                  borderwidth=0)
        browse_button.pack(pady=5)

# ------ END GUI Setup ------

## ------ UPDATE LOGS ------

    def update_log(self, message, status="info"):
        current_time = time.strftime("%H:%M:%S")
        self.conversion_table.add_row([message, status, current_time])
        
        log_content = str(self.conversion_table)
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.insert(tk.END, log_content)
        self.log_text.config(state=tk.DISABLED)
        self.log_text.see(tk.END)

    def _create_log_text_area(self):
        # Textfeld zur Anzeige des Logs (Konsolenausgabe)
        self.log_text = tk.Text(self.root, height=35, width=110, state=tk.DISABLED, bg="black", fg="#00FFFF",
                                insertbackground='#00FFFF',
                                font=('Consolas', 10))
        self.log_text.pack(padx=15, pady=10)

    def _initialize_conversion_table(self):
        # Initialisiere PrettyTable für die Konvertierungstabelle
        self.conversion_table = PrettyTable(["--Dateiname--", "--Status--", "--Zeit--"])
        self.conversion_table.align["--Dateiname--"] = "l"
        self.conversion_table.align["--Status--"] = "l"
        self.conversion_table.align["--Zeit--"] = "r"

    def _create_progress_bar(self):
        self.progress_bar = ttk.Progressbar(self.root, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=15, pady=5)
        self.progress_bar.pack_forget()

## ------ END UPDATE LOGS ------

##### ------ FILE SELECT + DRAG N DROP ------
    def drop(self, x, y, data_model, dx, dy):
        """Behandle das Drop-ereignis"""
        file_paths = [path for path in data_model.strip('{}').split() if path.lower().endswith('.pdf')]
        if file_paths:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, ", ".join(file_paths))
        else:
            messagebox.showwarning("Ungültige Datei", "Bitte wählen Sie nur PDF-Dateien aus.")
        
    def select_file(self):
        pdf_paths = filedialog.askopenfilenames(filetypes=[("PDF files", "*.pdf")])
        if pdf_paths:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, ", ".join(pdf_paths))
##### ------ END SELECT + DRAG N DROP ------

# ------ END OF INITIALISIERUNG ------

# -----------------------Konvertierung Beginn hier----------------------------------------
    async def start_conversion(self):
        try:
            self.progress_bar.pack(fill=tk.X, padx=15, pady=5)
            self.progress_var.set(0)

            pdf_paths = await self._get_pdf_paths()
            conversion_type = self.ask_conversoin_type()

            if not conversion_type: # Abbrechen ausgewählt
                self.update_log("Konvertierung abgebrochen.\n", "info\n")
                return

            total_files = len(pdf_paths)
            for i, pdf_path in enumerate(pdf_paths):
                output_dir = await self._create_output_directory(pdf_path, conversion_type)

                if conversion_type == 'docx':
                    await self.convert_pdf_to_docx(pdf_path, output_dir)
                
                progress = (i + 1) / total_files * 100
                self.progress_var.set(progress)
                await asyncio.sleep(0.1)
            self.update_log("Konvertierung abgeschlossen\n", "info\n")

        except Exception as e:
            self.update_log(f"Fehler während der Konvertierung: {str(e)}\n", "Error\n")
            messagebox.showerror("Fehler", f"Ein unerwarteter Fehler ist aufgetreten: {str(e)}")
        finally:
            self.progress_bar.pack_forget()

    async def _get_pdf_paths(self):
        pdf_paths = [path.strip() for path in self.file_entry.get().split(",") if path.strip()]
        
        valid_paths = []
        for pdf_path in pdf_paths:
            if os.path.exists(pdf_path) and pdf_path.lower().endswith('.pdf'):
                valid_paths.append(pdf_path)
            else:
                self.update_log(f"Ungültige Datei oder Pfad: {pdf_path}\n", "Error\n")
        
        return valid_paths

    async def _create_output_directory(self, pdf_path, conversion_type):
        output_dir = Path(pdf_path).parent / f"Konvertierte_{conversion_type.upper()}"

        try:
            os.makedirs(output_dir, exist_ok=True)
        except OSError as e:
            self.update_log(f"Fehler beim Erstellen des Ausgabeordners: {str(e)}\n", "Error\n")
        
        return output_dir    

    async def convert_pdf_to_docx(self, pdf_path, output_dir):
        try:
            pdf_filename = os.path.basename(pdf_path)
            docx_filename = os.path.splitext(pdf_filename)[0] + '.docx'
            docx_path = os.path.join(output_dir, docx_filename)
            
            if os.path.exists(docx_path):
                i = 1
                while os.path.exists(docx_path):    
                    docx_filename = f"{os.path.splitext(pdf_filename)[0]}({i}).docx"
                    docx_path = os.path.join(output_dir, docx_filename)
                    i += 1
            
            self.update_log(f"Starte Konvertierung in DOCX: {pdf_filename}\n", "fortschritt\n")
            
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            await self.optimize_docx(docx_path)
            self.update_log(f"Konvertierung in DOCX abgeschlossen: {docx_filename}\n", "erfolgreich\n")
            
            open_document = messagebox.askyesnocancel("Dokument öffnen", "Möchten Sie das konvertierte Dokument jetzt öffnen?")

            # Benutzer fragen ob Doc geöffnet werden soll
            if open_document:
                await self.open_document(docx_path)
            else:
                self.update_log(f"Dokument wurde nicht geöffnet: {docx_filename}\n", "info\n")
            
            return docx_path
        
        except Exception as e:
            self.update_log(f"Fehler bei der Konvertierung: {str(e)}\n", "Fehler\n")
            logging.error(f"Konvertierung fehlgeschlahgen: {str(e)}\n")
            return None
        
    async def optimize_docx(self, docx_path):
        try:
            self.update_log(f"Optimiere Dokument: {os.path.basename(docx_path)}\n", "bearbeiten\n")
            doc = Document(docx_path)
            for paragraph in doc.paragraphs:
                paragraph.text = ' '.join(paragraph.text.split())
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
            doc.save(docx_path)
            self.update_log(f"Dokument optimiert: {os.path.basename(docx_path)}\n", "erfolgreich\n")
        except Exception as e:
            self.update_log(f"Fehler bei der Optimierung der DOCX-Datei: {str(e)}\n", "Fehler\n")
            logging.error(f"DOCX Optimierung fehlgeschlagen: {str(e)}\n")

    async def open_document(self, docx_path):
        try:
            self.update_log(f"Öffne Dokument: {os.path.basename(docx_path)}\n", "info\n")
            os.startfile(docx_path)  
        except Exception as e:
            self.update_log(f"Fehler beim Öffnen des Dokuments: {str(e)}\n", "Fehler\n")
            logging.error(f"Fehler beim Öffnen des Dokuments: {str(e)}\n")

    def ask_conversoin_type(self):
        # Initialisiere Konversations Fenster
        conversion_window = tk.Toplevel(self.root)
        conversion_window.title("Dateityp auswählen")
        conversion_window.geometry("450x250")
        conversion_window.configure(bg='#1a1a1a')

        # Nachricht für User
        label = tk.Label(conversion_window, text="Wählen Sie den Dateityp für die Konvertierung:", bg='#1a1a1a', fg='#ffffff', font=('Helvetica', 12))
        label.pack(pady=10)

        conversion_type = tk.StringVar()

        # Buttons für Auswahl
        docx_button = tk.Button(conversion_window, text=".docx", command=lambda: self.set_conversion_type(conversion_window, conversion_type, 'docx'), bg='#333333', fg='#ffffff')
        docx_button.pack(pady=5)
        
        cancle_button = tk.Button(conversion_window, text="Abbrechen", command=lambda: self.set_conversion_type(conversion_window, conversion_type, None), bg='#333333', fg='#ffffff')
        cancle_button.pack(pady=5)

        self.root.wait_window(conversion_window) # Warte bis das Fenster geschlossen wird
        return conversion_type.get()

    def set_conversion_type(self, window, conversion_type_var, file_type):
        conversion_type_var.set(file_type)
        window.destroy()

    def run(self):
        self.root.mainloop()

# ---------------------------------------------------------------------------------------------------------------------------

def check_modules_installed():
    required_modules = ['tkinterdnd2', 'pdf2docx', 'python-docx', 'prettytable', 'docx']
    return all(importlib.util.find_spec(module) is not None for module in required_modules)

def save_installation_status(status):
    with open('installation_status.json', 'w') as f:
        json.dump({'installed': status}, f)

def load_installation_status():
    if os.path.exists('installatioin_status.json'):
        with open('installatioin_status.json', 'r') as f:
            return json.load(f).get('installed', False)
    return False

def check_tkinter():
    try:
        importlib.import_module('tkinter')
        print("tkinter ist bereits installiert.")
        return True
    except ImportError:
        print("tkinter ist nicht installiert. Bitte installieren Sie Python mit tkinter-UNterstützung.")
        return False

def install_required_packages():
    packages = [
        'tkinterdnd2',
        'pdf2docx',
        'python-docx',
        'prettytable'
    ]

    print("Beginne mit der Installation der erforderlichen Pakete...")

    for package in packages:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"Paket '{package}' erfolgreich installiert.")
        except subprocess.CalledProcessError as e:
            print(f"Fehler beim Installieren von '{package}': {e}")
            print("Bitte stellen Sie sicher, dass Sie die notwendigen Berechtigungen haben.")
            return False

    print("Alle erforderlichen Pakete wurden erfolgreich installiert.")
    return True

def ensure_dependencies():
    required_modules = ['tkinter', 'tkinterdnd2', 'pdf2docx']
    missing_modules = []
    for module in required_modules:
        try:
            __import__(module)
        except ImportError:
            missing_modules(module)

    if missing_modules:
        print(f"Warnung: Die folgenden Module fehlen: {', '.join(missing_modules)}")
        print("Bitte installieren Sie diese Module, bevor sie fortfahren")
        return False
    return True

def main():
    if load_installation_status() or check_modules_installed():
        # Alle Module sind installiert, starte die App direkt
        app = PDFMagicApp()
        app.run()
    else:
        # Module fehlen, führe die Installation durch
        if check_tkinter():
            if install_required_packages():
                if ensure_dependencies():
                    save_installation_status(True)
                    app = PDFMagicApp()
                    app.run()
                else:
                    print("Einige Abhängigkeiten fehlen noch. Bitte installieren Sie die fehlenden Module und versuchen Sie es erneut.")
            else:
                print("Installation einiger Pakete fehlgeschlagen. Bitte überprüfen Sie die Fehlermeldungen und versuchen Sie es erneut.")
        else:
            print("tkinter muss manuell installiert werden. Bitte installieren Sie Python mit tkinter-Unterstützung.")

# Treiber
if __name__ == "__main__":
    main()