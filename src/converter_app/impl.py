# Autor: Leon Gajtner
# Datum: 15.10.2024
# Projekt: PDF Magic Implementierungsdatei

import logging
from header import PDFConverterInterface, DropAreaInterface
from PyQt5.QtWidgets import QVBoxLayout, QLabel
from PyQt5.QtGui import QDragEnterEvent, QDropEvent
from PyQt5.QtCore import Qt 
from pdf2docx import Converter
import docx

# Logging einrichten
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PDFConverter(PDFConverterInterface):
    """
    Konkrete Implementierung des PDF-Konverters.
    """
    def run(self):
        """
        Hauptmethode zur Konvertierung von PDF-Dateien in DOCX.
        Verarbeitet jede PDF-Datei, aktualisiert den Fortschritt und gibt Statusmeldungen aus.
        """
        total_files = len(self.pdf_files)
        for index, pdf_file in enumerate(self.pdf_files, start=1):
            try:
                # Ausgabe-Dateiname generieren
                docx_file = pdf_file.rsplit('.', 1)[0] + '.docx'
                logger.info(f"Starte Konvertierung für Datei: {pdf_file}")

                # PDF in DOCX konvertieren
                cv = Converter(pdf_file)
                cv.convert(docx_file)
                cv.close()

                # Formatierung des konvertierten Dokuments anpassen
                doc = docx.Document(docx_file)
                for paragraph in doc.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = docx.shared.Pt(11)

                # Formatierung für Tabellen anpassen
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.style.font.name = 'Arial'
                                paragraph.style.font.size = docx.shared.Pt(11)

                # Formatiertes Dokument speichern
                doc.save(docx_file)
                self.update_log.emit(f"Erfolgreich konvertiert: {pdf_file}")
                logger.info(f"Erfolgreich konvertiert: {pdf_file}")
            except Exception as e:
                self.update_log.emit(f"Fehler bei der Konvertierung von {pdf_file}: {str(e)}")
                logger.error(f"Fehler bei der Konvertierung von {pdf_file}: {str(e)}")
            
            # Fortschritt aktualisieren
            progress = int((index / total_files) * 100)
            self.update_progress.emit(progress)

class DropArea(DropAreaInterface):
    """
    Konkrete Implementierung des Drop-Bereich-Widgets.
    """
    def __init__(self):
        """
        Initialisiere den Drop-Bereich mit einem Label.
        """
        super().__init__()
        layout = QVBoxLayout()
        self.label = QLabel("PDF-Dateien hier hineinziehen")
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label)
        self.setLayout(layout)

    def dragEnterEvent(self, event: QDragEnterEvent):
        """
        Handhabung von Drag-Eingabe-Ereignissen.
        Akzeptiere das Ereignis, wenn es URLs (mögliche Dateipfade) enthält.

        :param event: Das Drag-Eingabe-Ereignis
        """
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        """
        Handhabung von Drag-Bewegungs-Ereignissen.
        Erlaube die Kopieraktion, wenn das Ereignis URLs enthält.

        :param event: Das Drag-Bewegungs-Ereignis
        """
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        """
        Handhabung von Drop-Ereignissen.
        Verarbeite fallengelassene Dateien, wenn sie PDFs sind, und sende ein Signal mit der Dateiliste.

        :param event: Das Drop-Ereignis
        """
        files = [url.toLocalFile() for url in event.mimeData().urls() if url.toLocalFile().lower().endswith('.pdf')]
        if files:
            event.setDropAction(Qt.CopyAction)
            event.accept()
            self.files_dropped.emit(files)
        else:
            event.ignore()
