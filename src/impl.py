# Autor: Leon Gajtner
# Datum: 15.10.2024
# Projekt: PDF Magic impl file
from header import PDFConverterInterface
import logging
from PyQt5.QtCore import pyqtSignal
from pdf2docx import Converter
import docx
import os
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from PyPDF2 import PdfReader
import concurrent.futures

class PDFConverter(PDFConverterInterface):
    """
    Konkrete Implementierung des PDF-Konverters mit einheitlichem Logging.
    """
    update_progress = pyqtSignal(int)
    update_log = pyqtSignal(str)

    def __init__(self, pdf_files, save_dir):
        super().__init__(pdf_files) 
        self.pdf_files = pdf_files
        self.save_dir = save_dir
        self.logger = logging.getLogger(__name__)
        self.setup_logging()

    def setup_logging(self):
        """
        Richtet das Logging für die Klasse ein.
        """
        self.logger.setLevel(logging.INFO)
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

        # File handler
        file_handler = logging.FileHandler('pdf_converter.log')
        file_handler.setFormatter(formatter)
        self.logger.addHandler(file_handler)

        # Custom handler for GUI updates
        gui_handler = GUILogHandler(self.update_log)
        gui_handler.setFormatter(formatter)
        self.logger.addHandler(gui_handler)

    def log_info(self, message):
        """
        Protokolliert eine Informationsnachricht.
        """
        self.logger.info(message)

    def log_error(self, message):
        """
        Protokolliert eine Fehlermeldung.
        """
        self.logger.error(message)

    def run(self):
        """
        Hauptmethode zur Konvertierung von PDF-Dateien in DOCX.
        """
        total_files = len(self.pdf_files)
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.process_file, index, pdf_file) for index, pdf_file in enumerate(self.pdf_files, start=1)]
            for future in concurrent.futures.as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    self.log_error(f"Fehler bei der Verarbeitung einer Datei: {str(e)}")

    def process_file(self, index, pdf_file):
        try:
            # PDF in DOCX konvertieren
            docx_file = self.convert_pdf_to_docx(pdf_file)
            if docx_file:
                # Formatierung des konvertierten Dokuments anpassen
                self.format_docx(docx_file)

                # PDF in Bilder konvertieren
                self.convert_pdf_to_images(pdf_file)

                # Text aus PDF extrahieren
                self.extract_text_from_pdf(pdf_file)

                self.log_info(f"Erfolgreich verarbeitet: {pdf_file}")
            else:
                self.log_error(f"Fehler bei der Konvertierung von {pdf_file}")

            progress = int((index / len(self.pdf_files)) * 100)
            self.update_progress.emit(progress)
        except Exception as e:
            self.log_error(f"Fehler bei der Verarbeitung von {pdf_file}: {str(e)}")

    def convert_pdf_to_docx(self, pdf_file):
        """
        Konvertiere eine PDF-Datei in DOCX-Format.
        
        :param pdf_file: Pfad zur PDF-Datei.
        :return: Ausgegebener DOCX-Dateipfad oder None bei Fehlern.
        """
        try:
            output_dir = os.path.join(self.save_dir, 'docx')
            os.makedirs(output_dir, exist_ok=True)
            
            docx_file = os.path.join(output_dir, os.path.basename(pdf_file).rsplit('.', 1)[0] + '.docx')

            cv = Converter(pdf_file)
            cv.convert(docx_file)
            cv.close()
            
            return docx_file
        except Exception as e:
            self.log_error(f"Fehler bei der Konvertierung von PDF zu DOCX: {str(e)}")
            return None

    def format_docx(self, docx_file):
        """
        Passe die Formatierung des konvertierten DOCX-Dokuments an.
        """
        try:
            doc = docx.Document(docx_file)
            
            # Formatierung für alle Absätze anpassen
            for paragraph in doc.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = docx.shared.Pt(11)
            
            # Formatierung für Tabellen anpassen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style.font.name = 'Arial'
                            paragraph.style.font.size = docx.shared.Pt(11)
            
            # Formatiertes Dokument speichern
            doc.save(docx_file)
            
            self.log_info(f"Dokument formatiert und gespeichert: {docx_file}")
        except Exception as e:
            self.log_error(f"Fehler bei der Formatierung des Dokuments {docx_file}: {str(e)}")

    def convert_pdf_to_images(self, pdf_file):
        """
        Konvertiere eine PDF-Datei in Bilder (PNG).
        
        :param pdf_file: Pfad zur PDF-Datei.
        """
        try:
            output_dir = os.path.join(self.save_dir, 'images', os.path.basename(pdf_file).rsplit('.', 1)[0])
            os.makedirs(output_dir, exist_ok=True)
            
            images = convert_from_path(pdf_file)
            if not images:
                raise ValueError("Keine Bilder aus der PDF-Datei konvertiert.")
            
            for i, image in enumerate(images):
                image_file = os.path.join(output_dir, f"page_{i + 1}.png")
                image.save(image_file, 'PNG')
            
            self.log_info(f"PDF in Bilder konvertiert: {pdf_file}")
        except Exception as e:
            self.log_error(f"Fehler bei der Konvertierung von PDF zu Bildern: {str(e)}")

    def extract_text_from_pdf(self, pdf_file, use_ocr=False):
        """
        Extrahiere Text aus einer PDF-Datei. Optional kann OCR verwendet werden.
        
        :param pdf_file: Pfad zur PDF-Datei.
        :param use_ocr: Boolean, ob OCR für die Textextraktion verwendet werden soll.
        :return: Extrahierter Text als String.
        """
        try:
            extracted_text = ""
            
            # Versuche zuerst, den Text direkt mit PyPDF2 zu extrahieren
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                extracted_text += page.extract_text() or ""
            
            if not extracted_text and use_ocr:
                self.log_info("Kein Text gefunden. OCR wird verwendet.")
                images = convert_from_path(pdf_file)
                for image in images:
                    extracted_text += pytesseract.image_to_string(image)
            
            if not extracted_text:
                raise ValueError(f"Kein Text aus der PDF-Datei extrahiert: {pdf_file}")
            
            self.log_info(f"Text erfolgreich aus PDF extrahiert: {pdf_file}")
            return extracted_text
        except Exception as e:
            self.log_error(f"Fehler bei der Textextraktion aus der PDF: {str(e)}")
            return None

    def convert_from_file(self, file):
        """
        Konvertiere eine beliebige Datei (z. B. PDF oder Bild) in ein anderes Format.
        
        :param file: Pfad zur Eingabedatei.
        """
        try:
            if file.lower().endswith('.pdf'):
                self.convert_pdf_to_docx(file)
            elif file.lower().endswith((".png", ".jpg", ".jpeg")):
                output_pdf = os.path.join(self.save_dir, os.path.basename(file).rsplit('.', 1)[0] + '.pdf')
                image = Image.open(file)
                image.convert('RGB').save(output_pdf)
                self.log_info(f"Bild in PDF konvertiert: {file}")
            else:
                raise ValueError(f"Dateiformat wird nicht unterstützt: {file}")
        except Exception as e:
            self.log_error(f"Fehler bei der Dateikonvertierung: {str(e)}")

class GUILogHandler(logging.Handler):
    """
    Benutzerdefinierter Log-Handler zur Aktualisierung der GUI.
    """
    def __init__(self, signal):
        super().__init__()
        self.signal = signal

    def emit(self, record):
        log_entry = self.format(record)
        self.signal.emit(log_entry)