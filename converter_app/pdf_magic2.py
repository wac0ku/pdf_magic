import sys
import os
import logging
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QTextEdit, QFileDialog, QProgressBar, QLabel
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QDropEvent, QDragEnterEvent
from pdf2docx import Converter
import docx

class PDFConverter(QThread):
    update_progress = pyqtSignal(int)
    update_log = pyqtSignal(str)
    
    def __init__(self, pdf_files):
        super().__init__()
        self.pdf_files = pdf_files
    
    def run(self):
        total_files = len(self.pdf_files)
        for index, pdf_file in enumerate(self.pdf_files, start=1):
            try:
                docx_file = pdf_file.rsplit('.', 1)[0] + '.docx'
                cv = Converter(pdf_file)
                cv.convert(docx_file)
                cv.close()
                
                # Formatierung anpassen
                doc = docx.Document(docx_file)
                for paragraph in doc.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = docx.shared.Pt(11)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.style.font.name = 'Arial'
                                paragraph.style.font.size = docx.shared.Pt(11)
                
                doc.save(docx_file)
                
                self.update_log.emit(f"Erfolgreich konvertiert: {pdf_file}")
            except Exception as e:
                self.update_log.emit(f"Fehler bei der Konvertierung von {pdf_file}: {str(e)}")
            
            progress = int((index / total_files) * 100)
            self.update_progress.emit(progress)

class DropArea(QWidget):
    files_dropped = pyqtSignal(list)
    
    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        
        layout = QVBoxLayout()
        self.label = QLabel("PDF-Dateien hier hineinziehen")
        self.label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.label)
        
        self.setLayout(layout)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()
    
    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()
        else:
            event.ignore()
    
    def dropEvent(self, event: QDropEvent):
        files = [url.toLocalFile() for url in event.mimeData().urls() if url.toLocalFile().lower().endswith('.pdf')]
        if files:
            event.setDropAction(Qt.CopyAction)
            event.accept()
            self.files_dropped.emit(files)
        else:
            event.ignore()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF zu DOCX Konverter")
        self.setGeometry(100, 100, 600, 400)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        self.drop_area = DropArea()
        self.drop_area.setStyleSheet("""
            QWidget {
                border: 2px dashed #aaa;
                border-radius: 5px;
                background-color: #f0f0f0;
                min-height: 100px;
            }
            QWidget:hover {
                background-color: #e0e0e0;
            }
        """)
        self.drop_area.files_dropped.connect(self.add_files)
        layout.addWidget(self.drop_area)
        
        button_layout = QHBoxLayout()
        self.select_button = QPushButton("PDFs auswählen")
        self.select_button.clicked.connect(self.select_files)
        button_layout.addWidget(self.select_button)
        
        self.convert_button = QPushButton("Konvertieren")
        self.convert_button.clicked.connect(self.start_conversion)
        button_layout.addWidget(self.convert_button)
        
        layout.addLayout(button_layout)
        
        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)
        
        self.pdf_files = []
        
        # Logger konfigurieren
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        self.logger = logging.getLogger(__name__)
        
    def select_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "PDFs auswählen", "", "PDF Dateien (*.pdf)")
        self.add_files(files)
    
    def add_files(self, files):
        new_files = [f for f in files if f not in self.pdf_files]
        self.pdf_files.extend(new_files)
        self.log_text.append(f"{len(new_files)} neue PDF(s) hinzugefügt.")
        self.logger.info(f"{len(new_files)} neue PDF(s) zur Konvertierung hinzugefügt.")
        self.update_drop_area_label()
    
    def update_drop_area_label(self):
        if self.pdf_files:
            self.drop_area.label.setText(f"{len(self.pdf_files)} PDF(s) ausgewählt")
        else:
            self.drop_area.label.setText("PDF-Dateien hier hineinziehen")
    
    def start_conversion(self):
        if not self.pdf_files:
            self.log_text.append("Keine PDF-Dateien ausgewählt.")
            return
        
        self.converter = PDFConverter(self.pdf_files)
        self.converter.update_progress.connect(self.update_progress)
        self.converter.update_log.connect(self.update_log)
        self.converter.finished.connect(self.conversion_finished)
        self.converter.start()
        
        self.convert_button.setEnabled(False)
        self.select_button.setEnabled(False)
    
    def update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def update_log(self, message):
        self.log_text.append(message)
        self.logger.info(message)
    
    def conversion_finished(self):
        self.convert_button.setEnabled(True)
        self.select_button.setEnabled(True)
        self.pdf_files = []
        self.update_drop_area_label()
        self.log_text.append("Konvertierung abgeschlossen.")
        self.logger.info("Konvertierung abgeschlossen.")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())